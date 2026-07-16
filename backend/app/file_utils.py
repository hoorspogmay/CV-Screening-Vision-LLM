"""Resume text extraction utilities for PDF and DOCX files."""
from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(file_path)
    if suffix == ".docx":
        return _extract_docx_text(file_path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf_text(file_path: Path) -> str:
    try:
        text_parts: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        text = "\n".join(text_parts).strip()
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not read PDF: {exc}") from exc

    if not text:
        raise ValueError("No extractable text found in PDF (it may be a scanned image).")
    return text


def _extract_docx_text(file_path: Path) -> str:
    try:
        document = Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs).strip()
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not read DOCX: {exc}") from exc

    if not text:
        raise ValueError("No extractable text found in DOCX.")
    return text
